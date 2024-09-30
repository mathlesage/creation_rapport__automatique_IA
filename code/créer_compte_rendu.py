"""
Auteur : Mathéo Quatreboeufs
Couleur préférée : Orange

"""
from pathlib import Path
import sys
from openai import OpenAI
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
import config 

chemin_script = Path(sys.argv[0]).resolve()



_api_config = config.get_api_config()
_api_key = _api_config['api_key']
_api_endpoint = _api_config['api_endpoint']


# Ouvrir le fichier csv
def extract_data(chemin_csv: str) -> list:
    """Fonction qui extrait les données d'un fichier csv(seulement les bonnes commandes).

    Args:
        chemin_csv (str):chemin du fichier csv qui contient les données

    Returns:
        Une liste de dictionnaire chaque dictionnaire représente une ligne et une intervention.
    """
    #Sélection des colonnes souhaitées
    colonnes_a_lire = ["N° Demande","Libellé site","Ville","Motif de sollicitation","Message du client","Message du Client 2","Date/heure fin d'intervention","Problème réglé","Message au client"]
    #ouverture du document
    df = pd.read_csv(chemin_csv, encoding='cp1252', delimiter=';',usecols=colonnes_a_lire)
    L=[]
    for i in range(len(df)):
        #Chaque ligne est tranformée en dictionnaire
        a=dict(df.iloc[i])
        L.append(a)
    return L


# création du client



    
def write_prompt(donnee: dict) -> str:
    """Fonction qui écrit un prompte avec les arguments du dictionnaires.

    Args:
        dictionnaire contenant les informations sur l'intervention

    Returns:
        Retourn un prompt qui détaille les informations du techniciens
    """
# création du client

    client = OpenAI(api_key=_api_key, base_url=_api_endpoint)
    if("metier" in donnee):
        metier = donnee["metier"]
    else:
        metier = "techniciens"

    if(donnee["Problème réglé"] == "Non"):
        regler = "non reussite"
        prompttt = f"""
            Tu vas devoir rédiger tous le rapport d'intervention, je veux que tu sois soigné polie et que tu mettes les formes. 
            Voici les informations qu'on a : 
            -Tu es John Doe le techniciens en charge de l'intervention
            -le lieu est {donnee["Libellé site"]}
            -la date
            -la demande du client avant l'intervention est {donnee["Message du client"]}
            -la raison de l'echec donnée par toi le techniciens est {donnee["Message au client"]} 
            -l'intervention n'a pas été reussite
            -la date et l'heure de la fin de la mission est  {donnee["Date/heure fin d'intervention"]}
            -l'entreprise est Dal. 
            -Tu es John Doe le techniciens en charge de l'intervention
            Le rapport doit commencer par:
            -Bonjour,
            Je suis John Doe, technicien chez Dal, et je me permets de vous écrire suite à 
            Puis suivre l'ordre:
            -le lieux, 
            -la description de la demande, 
            -puis une description des interventions réalisées,  
            - tu peux conclure sur les résultats
            -Quand l'intervention n'est pas concluante tu exprimes des regrets.
            -Termine toujours par :"Nous restons à votre disposition pour toute information complémentaire. 
            N'hésitez pas à contacter le Centre de Relation Clients (CRC) au 0 800 80 93 00, disponible 24h/24 et 7j/7. 
            Nous tenons également à souligner l'importance de faire suivre votre matériel par un professionnel pour garantir sa qualité et sa durabilité".
            
            Mets toi à la place du techniciens John Doe qui a effectué l'intervention et parle à la première personne.
            Dit que la vérité et soit simple

            """
    else:
        regler = "reussite"

        prompttt = f"""
            Tu vas devoir rédiger tous le rapport d'intervention, je veux que tu sois soigné polie et que tu mettes les formes. 
            Voici les informations qu'on a : 
            -Tu es John Doe le techniciens en charge de l'intervention
            -le lieu est {donnee["Libellé site"]}
            -la date et l'heure de l'intervention sont {donnee["Date/heure fin d'intervention"]}
            -la demande du client avant l'intervention est {donnee["Message du client"]}
            -le message du techniciens après l'intervention est {donnee["Message au client"]} 
            -l'intervention a été réussite
            -l'entreprise est Dal. 
            Le rapport doit commencer par:
            -Bonjour,
            Je suis John Doe, technicien chez Dal, et je me permets de vous écrire suite à
            Puis suivre l'ordre:
            -le lieux, 
            -la date,
            -la description de la demande, 
            -puis une description des interventions réalisées,  
            - tu peux conclure sur les résultats
            -Termine par :"Nous restons à votre disposition pour toute information complémentaire. 
            N'hésitez pas à contacter le Centre de Relation Clients (CRC) au 0 800 80 93 00, disponible 24h/24 et 7j/7. 
            Nous tenons également à souligner l'importance de faire suivre votre matériel par un professionnel pour garantir sa qualité et sa durabilité".
            Mets toi à la place du techniciens John Doe qui a effectué l'intervention et parle à la première personne.
            Addresse toi directement au client.
            Dit que la vérité
            """
    
    appel_model = client.chat.completions.create(
    
    model="alfred-40b-1123", #alfred-vllm
    messages=[

        {"role": "system", "content" : f"Tu es John Doe un {metier} de chez Dal."},

        {"role":"user", "name":"fiche_de_poste",
        "content":prompttt}],
    temperature = 0.0)
    prompt = appel_model.choices[0].message.content
    return prompt, regler


def creation_document(date: str, 
                      num_demande: str, 
                      libelle: str, 
                      ville: str, 
                      sollicitation: str, 
                      compterendu: str, 
                      n: str,
                      message_du_client: str,
                      message_au_client: str,
                      regler: str
                      ):

    """Fonction qui avec les informations en argument fait un document world qui est le rapport de l'intervention.

    Args:
        Information relative à l'intervention en string.(chaque nom est très explicite)

    Returns:
        Return rien mais crée un document world représantant le rapport de l'intervention.
    """
    # Créer un nouveau document
    doc = Document()
    #date et demande 
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    row_cells = table.rows[0].cells
    row_cells[0].text = f"Date de l'intervention: {date}"
    #Ajout du logo
    paragraph = row_cells[1].paragraphs[0]
    run = paragraph.add_run()
    chemin_script = Path(sys.argv[0]).resolve()
    chemin_courant = chemin_script.parent.parent
    chemin_image=chemin_courant / 'image' / 'logo.png'
    run.add_picture(str(chemin_image), height=Inches(0.7))
    for cell in row_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.italic = True
                run.font.size = Pt(8)

    row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    #titre
    titre=doc.add_heading("Rapport d'intervention", 0)


    #information client
    soustitre=doc.add_heading("Informations client:",level=2)
    table = doc.add_table(rows=1, cols=2)
    row_cells = table.add_row().cells
    row_cells[0].text = "Libellé du site:"
    row_cells[1].text = f"{libelle}"


    row_cells = table.add_row().cells
    row_cells[0].text = "Ville:"
    row_cells[1].text = f"{ville}"

    row_cells = table.add_row().cells
    row_cells[0].text = "Motif de sollicitation:"
    row_cells[1].text = f"{sollicitation}"

    row_cells = table.add_row().cells
    row_cells[0].text = "Intervention "
    row_cells[1].text = f"{regler}"
    
    row_cells = table.add_row().cells
    row_cells[0].text = "Message du client:"
    row_cells[1].text = f"{message_du_client}"
    
    row_cells = table.add_row().cells
    row_cells[0].text = "Message au client:"
    row_cells[1].text = f"{message_au_client}"


    #Compte rendu de l'intervention
    soustitre=doc.add_heading("Compte rendu d'intervention:",level=2)
    paragrpahe_techniciens=doc.add_paragraph(f"{compterendu}\n\n")
    



    # Enregistrer le document
    chemin_script = Path(sys.argv[0]).resolve()
    chemin_courant = chemin_script.parent.parent
    dossier_data=chemin_courant / "rapport_world_speciale" / f'compte_rendu_{n}.docx'
    doc.save(dossier_data)



def main(url: str):
    chemin_courant = chemin_script.parent.parent
    dossier_data = chemin_courant / "rapport_world_speciale"
    if not(dossier_data.exists() and dossier_data.is_dir()):
        dossier_data.mkdir(parents=True, exist_ok=True)
    
    donnees = extract_data(url)
    for n,i in enumerate(donnees):

        compterendu,regler = write_prompt(i)
        creation_document(i["Date/heure fin d'intervention"],i["N° Demande"],i["Libellé site"],i["Ville"],i["Motif de sollicitation"],compterendu,n,i["Message du client"],i["Message au client"],regler)
        print(f"fichier {n+1}/{len(donnees)}")

if __name__ == "__main__":
    chemin_script = Path(sys.argv[0]).resolve()
    chemin_courant = chemin_script.parent.parent
    chemin_data = chemin_courant / "data\\tout.csv"
    
    main(chemin_data)



